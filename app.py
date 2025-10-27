import re
import os
import io
import tempfile
from flask import Flask, request, jsonify, send_file, send_from_directory
import fitz 
import firebase_admin
from firebase_admin import credentials, auth as firebase_auth
from dotenv import load_dotenv
from flask_cors import CORS
from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from authDecorator import firebase_required
from mcq_generator import generateMCQ 
from datetime import datetime
from narrative_generator import generateOpenEnded
from summarizing import summarize_topic
from werkzeug.utils import secure_filename
from eassy_generator import generate_essay_or_paragraph
from coverpage_generator import generate_coverpage
from assignment_generator import assignment_gen, markdown_to_plain_text
from study_plan import generate_study_plan
from topic_explanation import explain_topic
from paraphrasing import paraphrase_text
from pymongo import MongoClient


load_dotenv()
app = Flask(__name__)
CORS(app, origins=[
    "https://smart-study-assistant.vercel.app",
    "http://localhost:5173", 
    "http://127.0.0.1:5000"
], supports_credentials=True, methods=["GET", "POST", "OPTIONS"], allow_headers=["Authorization", "Content-Type"])

mongo_uri = os.getenv("MONGO_URI")

client = MongoClient(mongo_uri)
db = client["studyhelper"]
users_collection = db["users"]

cred = credentials.Certificate("firebase-service-account.json")
firebase_admin.initialize_app(cred)

def extractText(pdf_path):
    """Extract full text from a given PDF file."""
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text("text") for page in doc])
    return text.strip()

@app.route("/api/register-user", methods=["POST"])
def register_user():
    data = request.get_json()

    uid = data.get("uid")
    email = data.get("email")

    if not uid or not email:
        return jsonify({"error": "Missing uid or email"}), 400

    # Check if user already exists
    existing_user = users_collection.find_one({"uid": uid})
    if existing_user:
        return jsonify({"message": "User already exists"}), 200

    # Insert user
    users_collection.insert_one({
        "uid": uid,
        "email": email,
        "createdAt": datetime.utcnow()
    })

    return jsonify({"message": "User registered successfully"}), 201

def verify_token(request):
    auth_header = request.headers.get("Authorization")
    if not auth_header:
        return None
    token = auth_header.split("Bearer ")[1]
    #print("Authorization Header:", request.headers.get("Authorization"))
    try:
        decoded_token = firebase_auth.verify_id_token(token)
        return decoded_token["uid"]
    except Exception as e:
        print("Invalid token", e)
        return None

@app.route("/api/user", methods=["GET", "POST", "OPTIONS"])
def user_route():
    if request.method == "OPTIONS":
        return '', 200

    uid = verify_token(request)
    if not uid:
        return jsonify({"error": "Unauthorized"}), 401

    if request.method == "GET":
        user = users_collection.find_one({"uid": uid}, {"_id": 0})
        if not user:
            return jsonify({"message": "User not found"}), 404
        return jsonify(user), 200

    elif request.method == "POST":
        data = request.get_json()
        name = data.get("name", "").strip()
        university = data.get("university", "").strip()
        user_data = {"name": name, "university": university}
        users_collection.update_one({"uid": uid}, {"$set": user_data}, upsert=True)
        return jsonify({"message": "Profile updated"}), 200

@app.route("/api/user/texts", methods=["POST", "OPTIONS"])
def add_user_text():
    if request.method == "OPTIONS":
        return "", 200  # For CORS preflight
    print(">>> Headers:", dict(request.headers))
    print(">>> Raw body:", request.data)
    print(">>> Parsed JSON:", request.get_json(silent=True))

    # ✅ 1. Verify Firebase token and extract UID
    uid = verify_token(request)
    if not uid:
        return jsonify({"error": "Unauthorized"}), 401

    # ✅ 2. Parse and validate incoming JSON
    try:
        data = request.get_json(force=True)
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "Text is empty"}), 400
    except Exception as e:
        return jsonify({"error": "Invalid JSON"}), 400

    # ✅ 3. Prepare the entry
    entry = {
        "text": text,
        "createdAt": datetime.utcnow()
    }

    # ✅ 4. Push into the `texts` array of the user doc
    result = users_collection.update_one(
        {"uid": uid},
        {"$push": {"texts": entry}},
        upsert=True
    )

    # ✅ 5. Return success message
    return jsonify({
        "message": "Text saved successfully",
        "modifiedCount": result.modified_count
    }), 200

@app.route("/extract-text", methods=["POST"])
def extract_text_endpoint():
    """API endpoint to extract and return the text from an uploaded PDF."""
    
    if 'pdf' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    pdf_file = request.files['pdf']

    if pdf_file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    pdf_path = "temp.pdf"
    pdf_file.save(pdf_path)
    extracted_text = extractText(pdf_path)

    os.remove(pdf_path)

    if not extracted_text.strip():
        return jsonify({"error": "No text found in PDF"}), 400

    return jsonify({"extracted_text": extracted_text})

@app.route('/')
def home():
    return "<h1>Welcome to the PDF Question Generator API!</h1>"

@app.route("/generate-questions", methods=["POST"])
def generate_questions():
    """API endpoint to generate MCQs or Narrative Questions from PDF or Text."""
    
    question_type = request.form.get("question_type", "mcq")
    num_questions = int(request.form.get("num_questions", 10))
    level = request.form.get("level", "Undergraduate")
    difficulty = request.form.get("difficulty", "Average")

    extracted_text = ""

    if 'pdf' in request.files:
        pdf_file = request.files['pdf']
        pdf_path = "temp.pdf"
        pdf_file.save(pdf_path)
        extracted_text = extractText(pdf_path)
        os.remove(pdf_path)

    elif "text" in request.form:
        extracted_text = request.form.get("text", "")

    if not extracted_text.strip():
        return jsonify({"error": "No valid text provided."}), 400

    if question_type == "mcq":
        result = generateMCQ(extracted_text, num_questions, level, difficulty)
    else:
        result = generateOpenEnded(extracted_text, num_questions, level, difficulty)

    return jsonify(result)

@app.route("/summarize-pdf", methods=["POST"])
def summarize_pdf():
    """API endpoint to summarize text extracted from an uploaded PDF."""
    
    if 'pdf' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    pdf_file = request.files['pdf']
    
    if pdf_file.filename == "":
        return jsonify({"error": "Empty filename"}), 400
    
    pdf_path = "temp.pdf"
    pdf_file.save(pdf_path)
    extracted_text = extractText(pdf_path)
    
    if not extracted_text.strip():
        return jsonify({"error": "No text found in PDF"}), 400
    
    print("Processing PDF for summarization...")  # Debug log
    
    result = summarize_topic(extracted_text)
    
    os.remove(pdf_path)
    
    return jsonify({"summary": result})

@app.route("/explain-topic", methods=["POST"])
def explain_topic_api():
    """API endpoint to explain a given topic."""

    data = request.get_json()
    
    if not data or "topic" not in data:
        return jsonify({"error": "No topic provided"}), 400

    topic = data["topic"].strip()
    
    if not topic:
        return jsonify({"error": "Empty topic"}), 400
    
    print("Processing topic explanation...")  # Debug log

    result = explain_topic(topic)
    
    return jsonify({"explanation": result})

@app.route("/paraphrase", methods=["POST"])
def paraphrase():
    """API endpoint to paraphrase user input text."""
    
    data = request.get_json()
    
    if "text" not in data or not data["text"].strip():
        return jsonify({"error": "No text provided"}), 400

    print("Received text for paraphrasing...")  # Debug log

    result = paraphrase_text(data["text"])
    
    return jsonify({"paraphrased_text": result})


@app.route("/generate-study-plan", methods=["POST"])
def study_plan():
    data = request.json
    
    name = data.get("name")
    age = data.get("age")
    education_level = data.get("educationLevel")
    days_left = data.get("daysLeft")
    subjects = data.get("subjects")
    preferences = data.get("preferences")
    availability = data.get("availability")

    if not all([name, age, education_level, days_left, subjects, preferences, availability]):
        return jsonify({"error": "Missing required fields"}), 400

    print(f"Generating study plan for {name}...")  # Debug log
    
    study_plan = generate_study_plan(name, age, education_level, days_left, subjects, preferences, availability)
    
    return jsonify({"studyPlan": study_plan})


@app.route("/generate-essay", methods=["POST"])
def generate_essay():
    """API endpoint to generate an essay or paragraph based on a given topic."""
    
    data = request.json
    topic = data.get("topic", "").strip()
    essay_type = data.get("essay_type", "essay")  # "essay" or "paragraph"

    if not topic:
        return jsonify({"error": "No topic provided"}), 400

    print(f"Generating {essay_type} for topic: {topic}")  # Debugging log

    # Call the essay generation function
    result = generate_essay_or_paragraph(topic, essay_type)

    if "error" in result:
        return jsonify(result), 500

    return jsonify(result)


@app.route("/generate-lab-report", methods=["POST"])
def generate_Lab_Report():
    """API endpoint to generate a lab report from a given topic."""
    if not request.is_json:
        return jsonify({"error": "Invalid request. JSON data expected."}), 400
    
    data = request.get_json()
    topic = data.get("topic")
    if not topic:
        return jsonify({"error": "Missing topic"}), 400
    print(f"Received request for Lab Report on: {topic}")
    lab_report = assignment_gen(topic)
    if "error" in lab_report:
        return jsonify({"error": lab_report["error"]}), 500

    return jsonify(lab_report)

@app.route("/lab-report-docx", methods=["POST"])
def lab_report_docx():
    """API endpoint to generate a DOCX file for the lab report from a given topic."""
    if not request.is_json:
        return jsonify({"error": "Invalid request. JSON data expected."}), 400

    data = request.get_json()
    topic = data.get("topic")
    if not topic:
        return jsonify({"error": "Missing topic"}), 400

    print(f"Generating DOCX Lab Report for: {topic}")
    # Generate lab report using assignment_gen function
    lab_report_md = assignment_gen(topic)
    if isinstance(lab_report_md, dict) and "error" in lab_report_md:
        return jsonify({"error": lab_report_md["error"]}), 500

    # Convert Markdown to plain text
    plain_text = markdown_to_plain_text(lab_report_md)

    # Create a DOCX document from the plain text
    document = Document()
    for line in plain_text.split('\n'):
        document.add_paragraph(line)

    # Save the document into an in-memory bytes buffer
    file_buffer = io.BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)

    # Send the DOCX file as a downloadable attachment
    return send_file(
        file_buffer,
        as_attachment=True,
        download_name="lab_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/final-docx", methods=["POST"])
def final_docx():
    """
    API endpoint to generate a final DOCX that includes both a coverpage and a lab report.
    Expects JSON with:
      - topic: string
      - cover_data: object containing coverpage fields (optional; defaults are provided)
    """
    if not request.is_json:
        return jsonify({"error": "Invalid request. JSON data expected."}), 400

    data = request.get_json()
    topic = data.get("topic")
    if not topic:
        return jsonify({"error": "Missing topic"}), 400

    # Extract coverpage data from the request; use defaults if not provided.
    cover_data = data.get("cover_data", {})
    defaults = {
        "department": "Computer Science and Engineering",
        "course_code": "CSE-1102",
        "course_name": "Introduction to Programming Sessional",
        "assignment_name": "Lab Report Assignment",
        "date_of_submission": "01/01/2025",
        "submitted_by_name": "Student Name",
        "submitted_by_roll": "2103000",
        "submitted_by_section": "A",
        "submitted_by_series": "21",
        "submitted_to": "Dr. Teacher"
    }
    for key, value in defaults.items():
        if key not in cover_data or not cover_data[key]:
            cover_data[key] = value

    # 1. Generate Lab Report DOCX from markdown
    lab_report_md = assignment_gen(topic)
    if isinstance(lab_report_md, dict) and "error" in lab_report_md:
        return jsonify({"error": lab_report_md["error"]}), 500

    plain_text = markdown_to_plain_text(lab_report_md)
    lab_report_doc = Document()
    for line in plain_text.split('\n'):
        lab_report_doc.add_paragraph(line)

    # 2. Generate Coverpage DOCX using the coverpage template
    coverpage_doc_tpl = generate_coverpage(cover_data, template_path="coverpage_template.docx")
    # Save the rendered coverpage to a temporary file.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_cover:
        coverpage_doc_tpl.save(temp_cover.name)
    
    # 3. Save lab report document to a temporary file.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_lab:
        lab_report_doc.save(temp_lab.name)

    # 4. Merge the two documents using docxcompose (preserving each document's formatting)
    coverpage_doc = Document(temp_cover.name)
    lab_report_doc = Document(temp_lab.name)
    composer = Composer(coverpage_doc)
    composer.append(lab_report_doc)

    from docx.shared import Inches
    from docx.oxml.ns import qn

    merged_doc = composer.doc
    for section in merged_doc.sections[1:]:
        # Set margins to 1 inch on all sides
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Remove any background shading if present
        sectPr = section._sectPr
        shd = sectPr.find(qn('w:shd'))
        if shd is not None:
            sectPr.remove(shd)

    # Save the updated merged document to a temporary file.
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_final:
        composer.save(temp_final.name)
        temp_final.seek(0)
        final_data = temp_final.read()

    # Return the final merged document as a downloadable file.
    return send_file(
        io.BytesIO(final_data),
        as_attachment=True,
        download_name="final_lab_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    
    
@app.route('/submit', methods=['POST'])
def submit_data():
    data = request.get_json()
    print(data)
    return jsonify({"message": "Data received!", "data": data})


if __name__ == "__main__":
    from os import getenv
    port = int(getenv("PORT", 5000))  # Get the port from Railway
    app.run(host="0.0.0.0", port=port, debug=False)  # Bind to 0.0.0.0

